"""
Resume Parser
Extracts structured information from resume files (PDF, DOCX)
"""

import os
import re
from pathlib import Path
from typing import Dict, List, Any, Optional
from loguru import logger

import PyPDF2
import pdfplumber
from docx import Document

# For natural language processing
import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
from nltk.tag import pos_tag


class ResumeParser:
    """Parses resume files and extracts structured information"""
    
    def __init__(self, config_manager):
        """
        Initialize resume parser
        
        Args:
            config_manager: Configuration manager instance
        """
        self.config = config_manager
        self.resume_config = config_manager.get_resume_config()
        
        # Download required NLTK data
        self._setup_nltk()
        
        # Common skills and technologies patterns
        self.skill_patterns = self._load_skill_patterns()
        
    def _setup_nltk(self):
        """Download required NLTK data"""
        try:
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            logger.info("Downloading NLTK punkt tokenizer...")
            nltk.download('punkt', quiet=True)
            
        try:
            nltk.data.find('corpora/stopwords')
        except LookupError:
            logger.info("Downloading NLTK stopwords...")
            nltk.download('stopwords', quiet=True)
            
        try:
            nltk.data.find('taggers/averaged_perceptron_tagger')
        except LookupError:
            logger.info("Downloading NLTK POS tagger...")
            nltk.download('averaged_perceptron_tagger', quiet=True)
    
    def _load_skill_patterns(self) -> Dict[str, List[str]]:
        """Load common skill patterns for extraction"""
        return {
            'programming_languages': [
                'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'go', 'rust',
                'php', 'ruby', 'swift', 'kotlin', 'scala', 'r', 'matlab', 'sql'
            ],
            'web_technologies': [
                'html', 'css', 'react', 'angular', 'vue', 'node.js', 'express',
                'django', 'flask', 'spring', 'asp.net', 'laravel', 'rails'
            ],
            'databases': [
                'mysql', 'postgresql', 'mongodb', 'redis', 'sqlite', 'oracle',
                'sql server', 'elasticsearch', 'cassandra', 'dynamodb'
            ],
            'cloud_platforms': [
                'aws', 'azure', 'google cloud', 'gcp', 'docker', 'kubernetes',
                'terraform', 'ansible', 'jenkins', 'gitlab ci', 'github actions'
            ],
            'tools_frameworks': [
                'git', 'linux', 'ubuntu', 'bash', 'powershell', 'vim', 'vscode',
                'intellij', 'eclipse', 'postman', 'swagger', 'jira', 'confluence'
            ]
        }
    
    def parse_resume(self) -> Dict[str, Any]:
        """
        Parse resume and extract structured information
        
        Returns:
            Dictionary containing extracted resume information
        """
        resume_path = self.resume_config.get('file_path')
        if not resume_path:
            raise ValueError("Resume file path not configured")
        
        resume_path = Path(resume_path)
        if not resume_path.exists():
            raise FileNotFoundError(f"Resume file not found: {resume_path}")
        
        logger.info(f"Parsing resume: {resume_path}")
        
        # Extract text based on file format
        if resume_path.suffix.lower() == '.pdf':
            text = self._extract_pdf_text(resume_path)
        elif resume_path.suffix.lower() == '.docx':
            text = self._extract_docx_text(resume_path)
        else:
            raise ValueError(f"Unsupported file format: {resume_path.suffix}")
        
        # Extract structured information
        resume_data = self._extract_information(text)
        resume_data['raw_text'] = text
        resume_data['file_path'] = str(resume_path)
        
        logger.info(f"Resume parsing completed. Extracted {len(resume_data.get('skills', []))} skills")
        return resume_data
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        text = ""
        
        try:
            # Try with pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            except Exception as e2:
                logger.error(f"Failed to extract PDF text: {e2}")
                raise
        
        return text.strip()
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract DOCX text: {e}")
            raise
    
    def _extract_information(self, text: str) -> Dict[str, Any]:
        """Extract structured information from resume text"""
        
        resume_data = {
            'personal_info': self._extract_personal_info(text),
            'summary': self._extract_summary(text),
            'skills': self._extract_skills(text),
            'experience': self._extract_experience(text),
            'education': self._extract_education(text),
            'certifications': self._extract_certifications(text),
            'projects': self._extract_projects(text)
        }
        
        return resume_data
    
    def _extract_personal_info(self, text: str) -> Dict[str, str]:
        """Extract personal information (name, email, phone, location)"""
        info = {}
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            info['email'] = emails[0]
        
        # Extract phone number
        phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phones = re.findall(phone_pattern, text)
        if phones:
            info['phone'] = ''.join(phones[0]) if isinstance(phones[0], tuple) else phones[0]
        
        # Extract name (assume it's in the first few lines)
        lines = text.split('\n')[:5]
        for line in lines:
            line = line.strip()
            if line and not any(char in line for char in '@+()'):
                # Simple heuristic: if it's not too long and contains spaces, might be a name
                if 2 <= len(line.split()) <= 4 and len(line) < 50:
                    info['name'] = line
                    break
        
        return info
    
    def _extract_summary(self, text: str) -> str:
        """Extract professional summary or objective"""
        summary_keywords = ['summary', 'objective', 'profile', 'about']
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            if any(keyword in line.lower() for keyword in summary_keywords):
                # Get the next few lines as summary
                summary_lines = []
                for j in range(i + 1, min(i + 5, len(lines))):
                    if lines[j].strip():
                        summary_lines.append(lines[j].strip())
                    else:
                        break
                
                if summary_lines:
                    return ' '.join(summary_lines)
        
        return ""
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract technical skills"""
        text_lower = text.lower()
        skills = set()
        
        # Extract skills from all categories
        for category, skill_list in self.skill_patterns.items():
            for skill in skill_list:
                # Look for exact matches (with word boundaries)
                pattern = r'\b' + re.escape(skill.lower()) + r'\b'
                if re.search(pattern, text_lower):
                    skills.add(skill)
        
        # Look for skills section specifically
        skills_section = self._find_section(text, ['skills', 'technical skills', 'technologies'])
        if skills_section:
            # Extract comma-separated skills
            skill_items = re.split('[,;•\n]', skills_section)
            for item in skill_items:
                item = item.strip()
                if item and len(item) < 30:  # Filter out long descriptions
                    skills.add(item)
        
        return sorted(list(skills))
    
    def _extract_experience(self, text: str) -> List[Dict[str, str]]:
        """Extract work experience"""
        experience = []
        experience_section = self._find_section(text, ['experience', 'work experience', 'employment'])
        
        if not experience_section:
            return experience
        
        # Split by common job separators
        job_entries = re.split(r'\n(?=[A-Z][^a-z]*[A-Z])', experience_section)
        
        for entry in job_entries:
            if len(entry.strip()) < 20:  # Skip very short entries
                continue
                
            lines = [line.strip() for line in entry.split('\n') if line.strip()]
            if lines:
                job_info = {
                    'title': lines[0],
                    'description': ' '.join(lines[1:]) if len(lines) > 1 else ''
                }
                
                # Try to extract company and dates
                for line in lines[:3]:  # Check first few lines
                    # Look for date patterns
                    date_pattern = r'\b\d{4}\b'
                    if re.search(date_pattern, line):
                        job_info['dates'] = line
                        break
                
                experience.append(job_info)
        
        return experience
    
    def _extract_education(self, text: str) -> List[Dict[str, str]]:
        """Extract education information"""
        education = []
        education_section = self._find_section(text, ['education', 'academic background'])
        
        if not education_section:
            return education
        
        lines = [line.strip() for line in education_section.split('\n') if line.strip()]
        
        for line in lines:
            if any(degree in line.lower() for degree in ['bachelor', 'master', 'phd', 'degree', 'diploma']):
                education.append({
                    'degree': line,
                    'description': ''
                })
        
        return education
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications"""
        certifications = []
        cert_section = self._find_section(text, ['certification', 'certificates', 'licenses'])
        
        if cert_section:
            lines = [line.strip() for line in cert_section.split('\n') if line.strip()]
            certifications.extend(lines)
        
        return certifications
    
    def _extract_projects(self, text: str) -> List[Dict[str, str]]:
        """Extract project information"""
        projects = []
        project_section = self._find_section(text, ['projects', 'personal projects', 'portfolio'])
        
        if not project_section:
            return projects
        
        # Split by project indicators
        project_entries = re.split(r'\n(?=\s*[-•]\s*|\d+\.)', project_section)
        
        for entry in project_entries:
            if len(entry.strip()) < 10:
                continue
                
            lines = [line.strip() for line in entry.split('\n') if line.strip()]
            if lines:
                projects.append({
                    'name': lines[0],
                    'description': ' '.join(lines[1:]) if len(lines) > 1 else ''
                })
        
        return projects
    
    def _find_section(self, text: str, keywords: List[str]) -> Optional[str]:
        """Find a section in the resume text based on keywords"""
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            if any(keyword in line.lower() for keyword in keywords):
                # Extract section content
                section_lines = []
                for j in range(i + 1, len(lines)):
                    # Stop at next major section (all caps or common section headers)
                    next_line = lines[j].strip()
                    if (next_line.isupper() and len(next_line) > 3) or \
                       any(header in next_line.lower() for header in ['experience', 'education', 'skills', 'projects']):
                        break
                    section_lines.append(lines[j])
                
                return '\n'.join(section_lines)
        
        return None
